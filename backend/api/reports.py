"""
reports.py — Generates real .docx Word documents (python-docx), not markdown pretending to be
one. Two report types:
  - build_predict_report(): the Predict tab's results (single-role predictions + SHAP), as a
    clean narrative document with a chart, not a data dump.
  - build_chat_report(): a chat conversation, cleaned up into readable text (not a raw
    role/text transcript with markdown artifacts), plus a chart of any predictions the
    assistant actually ran as tools during that conversation, if any.

Both write to an in-memory BytesIO buffer (no temp files, no disk writes needed) and return
the buffer's bytes for the endpoint to stream back.
"""

from __future__ import annotations

import io
from datetime import datetime
from html import unescape
from typing import Optional
from xml.sax.saxutils import escape as _xml_escape

import matplotlib

matplotlib.use("Agg")  # headless — this process never opens a display
import matplotlib.pyplot as plt
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

from PIL import Image as PILImage
from reportlab.lib import colors as rl_colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    Image as RLImage,
    ListFlowable,
    ListItem,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table as RLTable,
    TableStyle,
)

# ZivaBasa's own gold/teal/red palette, matched to the frontend's Tailwind theme colors, so a
# chart embedded in a report doesn't look like it came from an unrelated tool.
GOLD = "#D4AF37"
TEAL = "#2FBF9F"
RED = "#D1495B"
INK = "#1F2430"

TASK_LABELS = {
    "employment": "Job & Automation Risk",
    "skills": "Employee Turnover Risk",
    "productivity": "AI Impact on Productivity",
    "skill_match": "Job and Skill Matching",
}

# Mirrors the frontend's fieldMeta.js labels (hand-duplicated on purpose — this backend module
# has no reason to import frontend JS, and a shared-taxonomy endpoint wasn't worth building for
# one label map). If you add a feature on either side, update both — same flagged tradeoff as
# skillMatchClient.js mirroring skill_matching.py's ALL_SKILLS.
FEATURE_LABELS = {
    "avg_salary_usd": "Salary",
    "ai_tool_maturity_score": "AI Tool Maturity",
    "task_repetition_level": "Task Repetition",
    "skill_complexity_score": "Skill Complexity",
    "training_hours_needed": "Training Hours Needed",
    "job_demand_index": "Job Demand",
    "percent_tasks_automatable": "% Tasks Automatable",
    "exposure_x_skill_complexity": "Exposure × Skill Complexity",
    "Age": "Age",
    "TrainingTimesLastYear": "Trainings Last Year",
    "YearsAtCompany": "Years at Company",
    "MonthlyIncome": "Monthly Income",
    "JobSatisfaction": "Job Satisfaction",
    "PerformanceRating": "Performance Rating",
    "training_intensity_index": "Training Intensity",
    "training_x_satisfaction": "Training × Satisfaction",
    "skill_gap_index": "Skill Gap",
    "seniority_years": "Seniority",
    "recent_training_hours": "Recent Training Hours",
    "performance_rating": "Performance Rating",
    "recent_ot_hours": "Recent Overtime Hours",
    "skill_overlap_count": "Matching Skills",
    "missing_skill_count": "Skill Gaps",
    "overlap_x_training": "Overlap × Training",
}


def _feature_label(name: str) -> str:
    return FEATURE_LABELS.get(name, name)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    """A real Word table — header row bold with light shading, borders via the built-in
    'Light Grid Accent 1' style rather than hand-rolling OOXML borders."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return table


def _add_title(doc: Document, text: str, subtitle: Optional[str] = None):
    heading = doc.add_heading(text, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        run.italic = True


def _add_disclaimer(doc: Document, text: str):
    doc.add_paragraph()  # spacer
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x8A, 0x90, 0x9C)
    run.italic = True


def _chart_to_stream(fig) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf


def _workflow_diagram(title: str, stages: list[str]) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(max(6.0, 1.65 * len(stages)), 2.0))
    ax.set_axis_off()
    ax.set_title(title, fontsize=11, pad=14)
    step_width = 0.78 / max(len(stages), 1)
    x_positions = [0.07 + i * (0.86 / max(len(stages), 1)) for i in range(len(stages))]
    for i, (x, stage) in enumerate(zip(x_positions, stages)):
        box = plt.Rectangle(
            (x, 0.32), 0.16, 0.28, facecolor="#F7F4EA", edgecolor=INK, linewidth=1.0
        )
        ax.add_patch(box)
        ax.text(
            x + 0.08,
            0.46,
            stage,
            ha="center",
            va="center",
            fontsize=8,
            color=INK,
            wrap=True,
        )
        if i < len(stages) - 1:
            ax.annotate(
                "",
                xy=(x + 0.16, 0.46),
                xytext=(x + 0.16 + step_width * 0.12, 0.46),
                arrowprops=dict(arrowstyle="->", color=TEAL, linewidth=1.4),
            )
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    fig.tight_layout()
    return _chart_to_stream(fig)


def _markdown_to_paragraph_runs(paragraph, node):
    if isinstance(node, NavigableString):
        text = unescape(str(node))
        if text:
            paragraph.add_run(text)
        return

    if not isinstance(node, Tag):
        return

    if node.name == "br":
        paragraph.add_run().add_break()
        return

    text = node.get_text(" ", strip=False)
    run = paragraph.add_run(text)
    if node.name in {"strong", "b"}:
        run.bold = True
    elif node.name in {"em", "i"}:
        run.italic = True
    elif node.name == "code":
        run.font.name = "Consolas"
        run.font.size = Pt(9)
    elif node.name == "a":
        run.underline = True


def _add_markdown_block(doc: Document, node):
    if isinstance(node, NavigableString):
        text = unescape(str(node)).strip()
        if text:
            doc.add_paragraph(text)
        return

    if not isinstance(node, Tag):
        return

    if node.name in {"div", "section", "article", "body"}:
        for child in node.children:
            _add_markdown_block(doc, child)
        return

    if node.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = min(int(node.name[1]), 4)
        heading = doc.add_heading(level=level)
        _markdown_to_paragraph_runs(heading, node)
        return

    if node.name == "p":
        paragraph = doc.add_paragraph()
        _markdown_to_paragraph_runs(paragraph, node)
        return

    if node.name in {"ul", "ol"}:
        style = "List Bullet" if node.name == "ul" else "List Number"
        for li in node.find_all("li", recursive=False):
            paragraph = doc.add_paragraph(style=style)
            _markdown_to_paragraph_runs(paragraph, li)
            nested_lists = [
                child
                for child in li.children
                if isinstance(child, Tag) and child.name in {"ul", "ol"}
            ]
            for nested in nested_lists:
                _add_markdown_block(doc, nested)
        return

    if node.name == "blockquote":
        paragraph = doc.add_paragraph(style="Intense Quote")
        _markdown_to_paragraph_runs(paragraph, node)
        return

    if node.name == "table":
        headers = []
        rows = []
        thead = node.find("thead")
        if thead:
            headers = [
                th.get_text(" ", strip=True)
                for th in thead.find_all("th", recursive=False)
            ]
        tbody = node.find("tbody") or node
        for tr in tbody.find_all("tr", recursive=False):
            cells = tr.find_all(["td", "th"], recursive=False)
            if not headers and cells and all(cell.name == "th" for cell in cells):
                headers = [cell.get_text(" ", strip=True) for cell in cells]
                continue
            if cells:
                rows.append([cell.get_text(" ", strip=True) for cell in cells])
        if headers and rows:
            _add_table(doc, headers, rows)
        return

    if node.name == "hr":
        doc.add_paragraph("—")
        return

    if node.name == "img":
        src = node.get("src")
        if src:
            doc.add_paragraph(f"[Image: {node.get('alt', src)}]")
        return

    paragraph = doc.add_paragraph()
    _markdown_to_paragraph_runs(paragraph, node)


def _add_markdown(doc: Document, text: str):
    html = markdown.markdown(text, extensions=["tables", "fenced_code", "extra"])
    soup = BeautifulSoup(f"<div>{html}</div>", "html.parser")
    for child in soup.div.children:
        _add_markdown_block(doc, child)


def _predictions_bar_chart(rows: list[dict]) -> io.BytesIO:
    """rows: [{label, value, kind: 'classification'|'regression'}]"""
    fig, ax = plt.subplots(figsize=(6, max(1.6, 0.6 * len(rows))))
    labels = [r["label"] for r in rows]
    values = [r["value"] for r in rows]
    colors = [
        TEAL
        if r["kind"] == "classification" and r["value"] < 0.5
        else (RED if r["kind"] == "classification" else GOLD)
        for r in rows
    ]
    ax.barh(labels, values, color=colors)
    ax.set_xlim(0, max(1.0, max(values) * 1.15) if values else 1.0)
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_xlabel("Score")
    fig.tight_layout()
    return _chart_to_stream(fig)


def _shap_chart(top_contributions: list[dict]) -> io.BytesIO:
    feats = [_feature_label(c["feature"]) for c in top_contributions][::-1]
    vals = [c["shap_value"] for c in top_contributions][::-1]
    colors = [TEAL if v >= 0 else RED for v in vals]
    fig, ax = plt.subplots(figsize=(6, max(1.6, 0.4 * len(feats))))
    ax.barh(feats, vals, color=colors)
    ax.axvline(0, color=INK, linewidth=0.8)
    ax.spines[["top", "right"]].set_visible(False)
    ax.set_title(f"What influenced this result", fontsize=10)
    fig.tight_layout()
    return _chart_to_stream(fig)


def build_predict_report(results: dict) -> bytes:
    """results: { task_name: { predict: {...}, explain: {...} | None } }, same shape the
    frontend's history entries already have — no reshaping needed at the call site."""
    doc = Document()
    _add_title(
        doc,
        "ZivaBasa Workforce Intelligence Report",
        f"Generated {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
    )

    _add_markdown(
        doc,
        """
This report summarizes predictions run on the Predict tab.

Each section below covers one prediction task in plain language, followed by the factors that most influenced that specific result.
        """.strip(),
    )

    doc.add_picture(
        _workflow_diagram(
            "Report flow",
            [
                "Input features",
                "Model prediction",
                "SHAP explanation",
                "Readable Word report",
            ],
        ),
        width=Inches(5.8),
    )

    chart_rows = []
    for task, r in results.items():
        if not r.get("predict"):
            continue
        p = r["predict"]
        if p["task_type"] == "classification":
            chart_rows.append(
                {
                    "label": TASK_LABELS.get(task, task),
                    "value": p["probability"],
                    "kind": "classification",
                }
            )
        else:
            chart_rows.append(
                {
                    "label": TASK_LABELS.get(task, task),
                    "value": p["raw_output"],
                    "kind": "regression",
                }
            )

    if chart_rows:
        doc.add_heading("At a glance", level=1)
        doc.add_picture(_predictions_bar_chart(chart_rows), width=Inches(5.5))
        _add_markdown(
            doc,
            """
| Prediction | Result | Score |
| --- | --- | --- |
"""
            + "\n".join(
                [
                    f"| {r['label']} | {'Flagged' if r['kind'] == 'classification' and r['value'] >= 0.5 else ('Not flagged' if r['kind'] == 'classification' else 'Standardized score')} | {r['value'] * 100:.1f}% |"
                    if r["kind"] == "classification"
                    else f"| {r['label']} | Standardized score | {r['value']:.3f} |"
                    for r in chart_rows
                ]
            ),
        )

    for task, r in results.items():
        label = TASK_LABELS.get(task, task)
        doc.add_heading(label, level=1)
        p = r.get("predict")
        if not p:
            doc.add_paragraph("Not run for this input.")
            continue

        if p["task_type"] == "classification":
            flagged = p["label"] == 1
            sentence = (
                f"This role was flagged for {label.lower()} "
                f"(estimated likelihood: {p['probability'] * 100:.1f}%)."
                if flagged
                else f"This role was not flagged for {label.lower()} "
                f"(estimated likelihood: {p['probability'] * 100:.1f}%)."
            )
        else:
            sentence = f"Predicted score: {p['raw_output']:.3f} (standardized — 0 is an average result, positive is above average, negative is below)."
        _add_markdown(doc, sentence)

        explain = r.get("explain")
        if explain and explain.get("top_contributions"):
            _add_markdown(doc, "**What influenced this result most:**")
            doc.add_picture(
                _shap_chart(explain["top_contributions"][:6]), width=Inches(5.5)
            )
            _add_markdown(
                doc,
                "Teal bars pushed the result up, red bars pushed it down. A longer bar means a bigger effect on this specific prediction.",
            )
            _add_table(
                doc,
                ["Factor", "Value", "Effect"],
                [
                    [
                        _feature_label(c["feature"]),
                        f"{c['value']:.3f}",
                        f"{'+' if c['shap_value'] >= 0 else ''}{c['shap_value']:.3f}",
                    ]
                    for c in explain["top_contributions"][:6]
                ],
            )

    _add_disclaimer(
        doc,
        "This is a prototype report built on Kaggle proxy / synthetic training data, not real "
        "company data. Treat every number here as illustrative, not a verified business "
        "finding. Explanations are local and associational (what mattered for this one "
        "prediction), not a claim about cause and effect.",
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _clean_chat_text(text: str) -> str:
    """Strips markdown clutter (##, **, bullet dashes) so the doc reads as prose, not a
    markdown file with the wrong extension — this is the whole point of this being a real
    Word doc instead of the old markdown-in-a-.md-file report."""
    import re

    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"^[-*]\s+", "• ", text, flags=re.MULTILINE)
    return text.strip()


def build_chat_report(messages: list[dict], tool_calls: list[dict]) -> bytes:
    """messages: [{role, text}]. tool_calls: [{name, args, result}] — every predict_task/
    explain_task call made anywhere in the conversation, in order, so the report can show what
    was actually predicted rather than just the free-text back-and-forth."""
    doc = Document()
    _add_title(
        doc,
        "ZivaBasa Chat Report",
        f"Generated {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
    )

    _add_markdown(
        doc,
        f"A record of this conversation with the ZivaBasa assistant ({len(messages)} message{'s' if len(messages) != 1 else ''}).",
    )

    doc.add_picture(
        _workflow_diagram(
            "Conversation flow",
            ["User message", "Model/tool response", "Readable Word report"],
        ),
        width=Inches(5.3),
    )

    predict_calls = [
        tc
        for tc in tool_calls
        if tc["name"] == "predict_task" and "error" not in tc.get("result", {})
    ]
    if predict_calls:
        doc.add_heading("Predictions made in this conversation", level=1)
        rows = []
        for tc in predict_calls:
            res = tc["result"]
            label = TASK_LABELS.get(res.get("task", ""), res.get("task", "?"))
            if res.get("task_type") == "classification":
                rows.append(
                    {
                        "label": label,
                        "value": res.get("probability", 0),
                        "kind": "classification",
                    }
                )
            else:
                rows.append(
                    {
                        "label": label,
                        "value": res.get("raw_output", 0),
                        "kind": "regression",
                    }
                )
        if rows:
            doc.add_picture(_predictions_bar_chart(rows), width=Inches(5.5))
            _add_table(
                doc,
                ["Prediction", "Result", "Score"],
                [
                    [
                        r["label"],
                        "Flagged"
                        if r["kind"] == "classification" and r["value"] >= 0.5
                        else (
                            "Not flagged"
                            if r["kind"] == "classification"
                            else "Standardized score"
                        ),
                        f"{r['value'] * 100:.1f}%"
                        if r["kind"] == "classification"
                        else f"{r['value']:.3f}",
                    ]
                    for r in rows
                ],
            )

    doc.add_heading("Conversation", level=1)
    for m in messages:
        speaker = "You" if m["role"] == "user" else "ZivaBasa Assistant"
        _add_markdown(doc, f"**{speaker}:** {_clean_chat_text(m.get('text', ''))}")

    _add_disclaimer(
        doc,
        "This is a record of an AI chat conversation from a prototype system trained on Kaggle "
        "proxy / synthetic data. Any predictions shown reflect the same limitations as the "
        "Predict tab — not verified business findings.",
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF export (reportlab) — same content and chart PNGs as the docx builders
# above, just walked into Platypus flowables instead of a Document object.
# Reportlab's Paragraph markup (<b>, <i>, <font>, <br/>) is a close enough
# match for our inline needs that we skip building an intermediate tree and
# convert the markdown HTML soup straight into flowables, mirroring the
# docx _add_markdown_block()/_markdown_to_paragraph_runs() pass above.
# ---------------------------------------------------------------------------

_PDF_STYLES = getSampleStyleSheet()
_PDF_STYLES.add(
    ParagraphStyle(
        name="ZBTitle",
        parent=_PDF_STYLES["Title"],
        alignment=TA_LEFT,
        textColor=rl_colors.HexColor(INK),
        fontSize=22,
        spaceAfter=4,
    )
)
_PDF_STYLES.add(
    ParagraphStyle(
        name="ZBSubtitle",
        parent=_PDF_STYLES["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=9,
        textColor=rl_colors.HexColor("#6B7280"),
        spaceAfter=10,
    )
)
_PDF_STYLES.add(
    ParagraphStyle(name="ZBH1", parent=_PDF_STYLES["Heading1"], textColor=rl_colors.HexColor(INK), spaceBefore=14, spaceAfter=8)
)
_PDF_STYLES.add(
    ParagraphStyle(name="ZBH2", parent=_PDF_STYLES["Heading2"], textColor=rl_colors.HexColor(INK), spaceBefore=10, spaceAfter=6)
)
_PDF_STYLES.add(
    ParagraphStyle(name="ZBH3", parent=_PDF_STYLES["Heading3"], textColor=rl_colors.HexColor(INK), spaceBefore=8, spaceAfter=4)
)
_PDF_STYLES.add(
    ParagraphStyle(name="ZBH4", parent=_PDF_STYLES["Heading4"], textColor=rl_colors.HexColor(INK), spaceBefore=6, spaceAfter=4)
)
_PDF_STYLES.add(
    ParagraphStyle(name="ZBBody", parent=_PDF_STYLES["Normal"], fontSize=10, leading=14, spaceAfter=6)
)
_PDF_STYLES.add(
    ParagraphStyle(
        name="ZBQuote",
        parent=_PDF_STYLES["Normal"],
        fontName="Helvetica-Oblique",
        leftIndent=16,
        textColor=rl_colors.HexColor("#4B5160"),
        spaceAfter=6,
    )
)
_PDF_STYLES.add(
    ParagraphStyle(
        name="ZBDisclaimer",
        parent=_PDF_STYLES["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=8,
        textColor=rl_colors.HexColor("#8A909C"),
    )
)


def _pdf_image(stream: io.BytesIO, width: float) -> RLImage:
    stream.seek(0)
    pil_img = PILImage.open(stream)
    aspect = pil_img.height / pil_img.width
    stream.seek(0)
    return RLImage(stream, width=width, height=width * aspect)


def _pdf_table(headers: list[str], rows: list[list[str]]) -> RLTable:
    table = RLTable([headers] + [[str(v) for v in row] for row in rows], hAlign="LEFT", repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), rl_colors.HexColor("#F2EFE4")),
                ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.HexColor(INK)),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.HexColor("#D9D9D9")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [rl_colors.white, rl_colors.HexColor("#FAFAF7")]),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def _pdf_inline_markup(node) -> str:
    """Renders a bs4 node to reportlab's Paragraph mini-markup (a restricted, non-extensible
    subset of HTML: <b>, <i>, <u>, <br/>, <font>) instead of docx runs."""
    if isinstance(node, NavigableString):
        return _xml_escape(unescape(str(node)))
    if not isinstance(node, Tag):
        return ""
    if node.name == "br":
        return "<br/>"
    inner = "".join(_pdf_inline_markup(child) for child in node.children)
    if node.name in {"strong", "b"}:
        return f"<b>{inner}</b>"
    if node.name in {"em", "i"}:
        return f"<i>{inner}</i>"
    if node.name == "code":
        return f'<font face="Courier" size="9">{inner}</font>'
    if node.name == "a":
        return f"<u>{inner}</u>"
    return inner


def _pdf_markdown_block(story: list, node):
    if isinstance(node, NavigableString):
        text = unescape(str(node)).strip()
        if text:
            story.append(Paragraph(_xml_escape(text), _PDF_STYLES["ZBBody"]))
        return

    if not isinstance(node, Tag):
        return

    if node.name in {"div", "section", "article", "body"}:
        for child in node.children:
            _pdf_markdown_block(story, child)
        return

    if node.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
        level = min(int(node.name[1]), 4)
        story.append(Paragraph(_pdf_inline_markup(node), _PDF_STYLES[f"ZBH{level}"]))
        return

    if node.name == "p":
        story.append(Paragraph(_pdf_inline_markup(node), _PDF_STYLES["ZBBody"]))
        return

    if node.name in {"ul", "ol"}:
        items = []
        for li in node.find_all("li", recursive=False):
            direct_markup = "".join(
                _pdf_inline_markup(child)
                for child in li.children
                if not (isinstance(child, Tag) and child.name in {"ul", "ol"})
            )
            items.append(ListItem(Paragraph(direct_markup, _PDF_STYLES["ZBBody"])))
        story.append(
            ListFlowable(
                items,
                bulletType="bullet" if node.name == "ul" else "1",
                leftIndent=18,
            )
        )
        nested_lists = [
            child
            for li in node.find_all("li", recursive=False)
            for child in li.children
            if isinstance(child, Tag) and child.name in {"ul", "ol"}
        ]
        for nested in nested_lists:
            _pdf_markdown_block(story, nested)
        return

    if node.name == "blockquote":
        story.append(Paragraph(_pdf_inline_markup(node), _PDF_STYLES["ZBQuote"]))
        return

    if node.name == "table":
        headers = []
        rows = []
        thead = node.find("thead")
        if thead:
            headers = [
                th.get_text(" ", strip=True)
                for th in thead.find_all("th", recursive=False)
            ]
        tbody = node.find("tbody") or node
        for tr in tbody.find_all("tr", recursive=False):
            cells = tr.find_all(["td", "th"], recursive=False)
            if not headers and cells and all(cell.name == "th" for cell in cells):
                headers = [cell.get_text(" ", strip=True) for cell in cells]
                continue
            if cells:
                rows.append([cell.get_text(" ", strip=True) for cell in cells])
        if headers and rows:
            story.append(_pdf_table(headers, rows))
        return

    if node.name == "hr":
        story.append(
            HRFlowable(width="100%", color=rl_colors.HexColor("#D9D9D9"), thickness=0.75, spaceBefore=6, spaceAfter=6)
        )
        return

    if node.name == "img":
        src = node.get("src")
        if src:
            story.append(Paragraph(f"[Image: {_xml_escape(node.get('alt', src))}]", _PDF_STYLES["ZBBody"]))
        return

    story.append(Paragraph(_pdf_inline_markup(node), _PDF_STYLES["ZBBody"]))


def _add_markdown_pdf(story: list, text: str):
    html = markdown.markdown(text, extensions=["tables", "fenced_code", "extra"])
    soup = BeautifulSoup(f"<div>{html}</div>", "html.parser")
    for child in soup.div.children:
        _pdf_markdown_block(story, child)


def _pdf_title_block(story: list, text: str, subtitle: Optional[str] = None):
    story.append(Paragraph(_xml_escape(text), _PDF_STYLES["ZBTitle"]))
    if subtitle:
        story.append(Paragraph(_xml_escape(subtitle), _PDF_STYLES["ZBSubtitle"]))


def _pdf_disclaimer_block(story: list, text: str):
    story.append(Spacer(1, 14))
    story.append(Paragraph(_xml_escape(text), _PDF_STYLES["ZBDisclaimer"]))


def build_predict_report_pdf(results: dict) -> bytes:
    """Same input contract and content as build_predict_report() — a PDF sibling of the docx
    report, sharing the same matplotlib chart builders."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="ZivaBasa Workforce Intelligence Report",
    )
    story: list = []
    _pdf_title_block(
        story,
        "ZivaBasa Workforce Intelligence Report",
        f"Generated {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
    )

    _add_markdown_pdf(
        story,
        """
This report summarizes predictions run on the Predict tab.

Each section below covers one prediction task in plain language, followed by the factors that most influenced that specific result.
        """.strip(),
    )

    story.append(
        _pdf_image(
            _workflow_diagram(
                "Report flow",
                [
                    "Input features",
                    "Model prediction",
                    "SHAP explanation",
                    "Readable PDF report",
                ],
            ),
            width=5.8 * inch,
        )
    )
    story.append(Spacer(1, 8))

    chart_rows = []
    for task, r in results.items():
        if not r.get("predict"):
            continue
        p = r["predict"]
        if p["task_type"] == "classification":
            chart_rows.append(
                {
                    "label": TASK_LABELS.get(task, task),
                    "value": p["probability"],
                    "kind": "classification",
                }
            )
        else:
            chart_rows.append(
                {
                    "label": TASK_LABELS.get(task, task),
                    "value": p["raw_output"],
                    "kind": "regression",
                }
            )

    if chart_rows:
        story.append(Paragraph("At a glance", _PDF_STYLES["ZBH1"]))
        story.append(_pdf_image(_predictions_bar_chart(chart_rows), width=5.5 * inch))
        story.append(Spacer(1, 6))
        _add_markdown_pdf(
            story,
            """
| Prediction | Result | Score |
| --- | --- | --- |
"""
            + "\n".join(
                [
                    f"| {r['label']} | {'Flagged' if r['kind'] == 'classification' and r['value'] >= 0.5 else ('Not flagged' if r['kind'] == 'classification' else 'Standardized score')} | {r['value'] * 100:.1f}% |"
                    if r["kind"] == "classification"
                    else f"| {r['label']} | Standardized score | {r['value']:.3f} |"
                    for r in chart_rows
                ]
            ),
        )

    for task, r in results.items():
        label = TASK_LABELS.get(task, task)
        story.append(Paragraph(label, _PDF_STYLES["ZBH1"]))
        p = r.get("predict")
        if not p:
            story.append(Paragraph("Not run for this input.", _PDF_STYLES["ZBBody"]))
            continue

        if p["task_type"] == "classification":
            flagged = p["label"] == 1
            sentence = (
                f"This role was flagged for {label.lower()} "
                f"(estimated likelihood: {p['probability'] * 100:.1f}%)."
                if flagged
                else f"This role was not flagged for {label.lower()} "
                f"(estimated likelihood: {p['probability'] * 100:.1f}%)."
            )
        else:
            sentence = f"Predicted score: {p['raw_output']:.3f} (standardized — 0 is an average result, positive is above average, negative is below)."
        _add_markdown_pdf(story, sentence)

        explain = r.get("explain")
        if explain and explain.get("top_contributions"):
            _add_markdown_pdf(story, "**What influenced this result most:**")
            story.append(_pdf_image(_shap_chart(explain["top_contributions"][:6]), width=5.5 * inch))
            story.append(Spacer(1, 4))
            _add_markdown_pdf(
                story,
                "Teal bars pushed the result up, red bars pushed it down. A longer bar means a bigger effect on this specific prediction.",
            )
            story.append(
                _pdf_table(
                    ["Factor", "Value", "Effect"],
                    [
                        [
                            _feature_label(c["feature"]),
                            f"{c['value']:.3f}",
                            f"{'+' if c['shap_value'] >= 0 else ''}{c['shap_value']:.3f}",
                        ]
                        for c in explain["top_contributions"][:6]
                    ],
                )
            )
            story.append(Spacer(1, 6))

    _pdf_disclaimer_block(
        story,
        "This is a prototype report built on Kaggle proxy / synthetic training data, not real "
        "company data. Treat every number here as illustrative, not a verified business "
        "finding. Explanations are local and associational (what mattered for this one "
        "prediction), not a claim about cause and effect.",
    )

    doc.build(story)
    return buf.getvalue()


def build_chat_report_pdf(messages: list[dict], tool_calls: list[dict]) -> bytes:
    """Same input contract and content as build_chat_report() — a PDF sibling of the docx
    report, sharing the same matplotlib chart builders and _clean_chat_text() cleanup."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.8 * inch,
        title="ZivaBasa Chat Report",
    )
    story: list = []
    _pdf_title_block(
        story,
        "ZivaBasa Chat Report",
        f"Generated {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
    )

    _add_markdown_pdf(
        story,
        f"A record of this conversation with the ZivaBasa assistant ({len(messages)} message{'s' if len(messages) != 1 else ''}).",
    )

    story.append(
        _pdf_image(
            _workflow_diagram(
                "Conversation flow",
                ["User message", "Model/tool response", "Readable PDF report"],
            ),
            width=5.3 * inch,
        )
    )
    story.append(Spacer(1, 8))

    predict_calls = [
        tc
        for tc in tool_calls
        if tc["name"] == "predict_task" and "error" not in tc.get("result", {})
    ]
    if predict_calls:
        story.append(Paragraph("Predictions made in this conversation", _PDF_STYLES["ZBH1"]))
        rows = []
        for tc in predict_calls:
            res = tc["result"]
            label = TASK_LABELS.get(res.get("task", ""), res.get("task", "?"))
            if res.get("task_type") == "classification":
                rows.append(
                    {
                        "label": label,
                        "value": res.get("probability", 0),
                        "kind": "classification",
                    }
                )
            else:
                rows.append(
                    {
                        "label": label,
                        "value": res.get("raw_output", 0),
                        "kind": "regression",
                    }
                )
        if rows:
            story.append(_pdf_image(_predictions_bar_chart(rows), width=5.5 * inch))
            story.append(Spacer(1, 6))
            story.append(
                _pdf_table(
                    ["Prediction", "Result", "Score"],
                    [
                        [
                            r["label"],
                            "Flagged"
                            if r["kind"] == "classification" and r["value"] >= 0.5
                            else (
                                "Not flagged"
                                if r["kind"] == "classification"
                                else "Standardized score"
                            ),
                            f"{r['value'] * 100:.1f}%"
                            if r["kind"] == "classification"
                            else f"{r['value']:.3f}",
                        ]
                        for r in rows
                    ],
                )
            )
            story.append(Spacer(1, 6))

    story.append(Paragraph("Conversation", _PDF_STYLES["ZBH1"]))
    for m in messages:
        speaker = "You" if m["role"] == "user" else "ZivaBasa Assistant"
        _add_markdown_pdf(story, f"**{speaker}:** {_clean_chat_text(m.get('text', ''))}")

    _pdf_disclaimer_block(
        story,
        "This is a record of an AI chat conversation from a prototype system trained on Kaggle "
        "proxy / synthetic data. Any predictions shown reflect the same limitations as the "
        "Predict tab — not verified business findings.",
    )

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Excel export (openpyxl) — a data-first sibling of the docx/PDF reports: one
# row per prediction/factor for stakeholders who want to filter/pivot the
# numbers rather than read a narrative document.
# ---------------------------------------------------------------------------

_XL_HEADER_FILL = PatternFill(start_color="F2EFE4", end_color="F2EFE4", fill_type="solid")
_XL_HEADER_FONT = Font(bold=True, color=INK.lstrip("#"))


def _xl_write_table(ws, start_row: int, headers: list[str], rows: list[tuple]) -> int:
    """Writes a header row (styled) + data rows starting at start_row; returns the next
    free row so callers can stack multiple tables on one sheet."""
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=start_row, column=col, value=h)
        cell.font = _XL_HEADER_FONT
        cell.fill = _XL_HEADER_FILL
    for r, row in enumerate(rows, start=start_row + 1):
        for col, val in enumerate(row, start=1):
            ws.cell(row=r, column=col, value=val)
    for col in range(1, len(headers) + 1):
        ws.column_dimensions[get_column_letter(col)].width = 24
    return start_row + 1 + len(rows)


def _xl_embed_chart(ws, anchor_cell: str, stream: io.BytesIO, width_px: int = 440):
    """Embeds one of the shared matplotlib PNG charts (same builders the docx/PDF reports
    use) into a worksheet at anchor_cell, scaled to width_px."""
    stream.seek(0)
    pil_img = PILImage.open(stream)
    scale = width_px / pil_img.width
    stream.seek(0)
    img = XLImage(stream)
    img.width = width_px
    img.height = int(pil_img.height * scale)
    ws.add_image(img, anchor_cell)


def build_predict_report_xlsx(results: dict) -> bytes:
    """Same input contract as build_predict_report()/_pdf() — a Summary sheet (table +
    chart) plus one sheet per predicted task with its top SHAP contributions."""
    wb = Workbook()
    summary = wb.active
    summary.title = "Summary"
    summary["A1"] = "ZivaBasa Workforce Intelligence Report"
    summary["A1"].font = Font(bold=True, size=14)
    summary["A2"] = f"Generated {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    summary["A2"].font = Font(italic=True, size=9, color="6B7280")

    predicted = {task: r for task, r in results.items() if r.get("predict")}
    table_rows = []
    chart_rows = []
    for task, r in predicted.items():
        p = r["predict"]
        label = TASK_LABELS.get(task, task)
        if p["task_type"] == "classification":
            table_rows.append((label, "Flagged" if p["label"] == 1 else "Not flagged", f"{p['probability'] * 100:.1f}%"))
        else:
            table_rows.append((label, "Standardized score", f"{p['raw_output']:.3f}"))
        chart_rows.append({
            "label": label,
            "value": p["probability"] if p["task_type"] == "classification" else p["raw_output"],
            "kind": p["task_type"],
        })

    if table_rows:
        _xl_write_table(summary, 4, ["Prediction", "Result", "Score"], table_rows)
        _xl_embed_chart(summary, "E4", _predictions_bar_chart(chart_rows))
    else:
        summary["A4"] = "No predictions were run for this input."

    for task, r in predicted.items():
        label = TASK_LABELS.get(task, task)
        p = r["predict"]
        explain = r.get("explain")
        ws = wb.create_sheet(label[:31])  # Excel sheet-name length limit
        ws["A1"] = label
        ws["A1"].font = Font(bold=True, size=12)
        ws["A2"] = (
            f"Estimated likelihood: {p['probability'] * 100:.1f}% ({'flagged' if p['label'] == 1 else 'not flagged'})"
            if p["task_type"] == "classification"
            else f"Predicted score: {p['raw_output']:.3f}"
        )

        if explain and explain.get("top_contributions"):
            rows = [
                (_feature_label(c["feature"]), c["value"], c["shap_value"])
                for c in explain["top_contributions"][:6]
            ]
            _xl_write_table(ws, 4, ["Factor", "Value", "SHAP effect"], rows)
            _xl_embed_chart(ws, "E4", _shap_chart(explain["top_contributions"][:6]))

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def build_chat_report_xlsx(messages: list[dict], tool_calls: list[dict]) -> bytes:
    """Spreadsheet sibling of build_chat_report()/_pdf(): a Predictions sheet (any
    predict_task tool calls made during the conversation) + a Conversation transcript sheet."""
    wb = Workbook()
    summary = wb.active
    summary.title = "Predictions"
    summary["A1"] = "ZivaBasa Chat Report"
    summary["A1"].font = Font(bold=True, size=14)
    summary["A2"] = f"Generated {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    summary["A2"].font = Font(italic=True, size=9, color="6B7280")

    predict_calls = [
        tc for tc in tool_calls
        if tc["name"] == "predict_task" and "error" not in tc.get("result", {})
    ]
    rows = []
    for tc in predict_calls:
        res = tc["result"]
        label = TASK_LABELS.get(res.get("task", ""), res.get("task", "?"))
        if res.get("task_type") == "classification":
            rows.append((label, "Flagged" if res.get("probability", 0) >= 0.5 else "Not flagged", f"{res.get('probability', 0) * 100:.1f}%"))
        else:
            rows.append((label, "Standardized score", f"{res.get('raw_output', 0):.3f}"))

    if rows:
        _xl_write_table(summary, 4, ["Prediction", "Result", "Score"], rows)
    else:
        summary["A4"] = "No predictions were made in this conversation."

    transcript = wb.create_sheet("Conversation")
    transcript["A1"] = "Speaker"
    transcript["B1"] = "Message"
    transcript["A1"].font = Font(bold=True)
    transcript["B1"].font = Font(bold=True)
    for i, m in enumerate(messages, start=2):
        speaker = "You" if m["role"] == "user" else "ZivaBasa Assistant"
        transcript.cell(row=i, column=1, value=speaker)
        transcript.cell(row=i, column=2, value=_clean_chat_text(m.get("text", "")))
    transcript.column_dimensions["A"].width = 20
    transcript.column_dimensions["B"].width = 100

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()
